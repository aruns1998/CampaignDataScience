import pandas as pd
from docx import Document
from docx.shared import Pt

# Load CSV files
election = pd.read_csv('data/election_history.csv')
posts = pd.read_csv('data/sentiment_posts.csv')
features = pd.read_csv('data/features.csv')

# Create Word document
doc = Document()
doc.add_heading('Data-Driven Political Campaign & Sentiment Analysis', 0)
doc.add_heading('Thiruvidaimaruthur Constituency (2006–2024)', level=1)

# ----------------------------
# 1. Constituency Overview
# ----------------------------
doc.add_heading('1. Constituency Overview', level=2)
doc.add_paragraph(
    'Location: Thiruvidaimaruthur, Tamil Nadu\n'
    'Demographics: Population details, SC reservation, voter profile'
)

# ----------------------------
# 2. Election History
# ----------------------------
doc.add_heading('2. Election History (2006–2021)', level=2)
table = doc.add_table(rows=1, cols=len(election.columns))
table.style = 'Table Grid'

# Add headers
hdr_cells = table.rows[0].cells
for i, col in enumerate(election.columns):
    hdr_cells[i].text = col

# Add rows
for _, row in election.iterrows():
    row_cells = table.add_row().cells
    for i, col in enumerate(election.columns):
        row_cells[i].text = str(row[col])

# ----------------------------
# 3. Social Media Sentiment Summary
# ----------------------------
doc.add_heading('3. Social Media Sentiment Summary', level=2)
if 'sentiment_score' not in posts.columns:
    from textblob import TextBlob
    posts['sentiment_score'] = posts['post_text'].apply(lambda x: TextBlob(str(x)).sentiment.polarity)

sentiment_summary = posts.groupby('booth_id')['sentiment_score'].mean().reset_index()
for _, row in sentiment_summary.iterrows():
    doc.add_paragraph(f"Booth {row['booth_id']} - Avg Sentiment Score: {row['sentiment_score']:.2f}")

# ----------------------------
# 4. Booth-level Features
# ----------------------------
doc.add_heading('4. Booth-level Features', level=2)
table = doc.add_table(rows=1, cols=len(features.columns))
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
for i, col in enumerate(features.columns):
    hdr_cells[i].text = col

for _, row in features.iterrows():
    row_cells = table.add_row().cells
    for i, col in enumerate(features.columns):
        row_cells[i].text = str(row[col])

# ----------------------------
# 5. ML Predictions (Temp Winning Probability)
# ----------------------------
doc.add_heading('5. Predicted Winning Probability', level=2)
for _, row in features.iterrows():
    doc.add_paragraph(f"Booth {row['Booth_ID']} - Winning Probability: {row['temp_winning_prob']:.2f}")

# ----------------------------
# Save the document
# ----------------------------
doc.save('output/campaign_analysis.docx')
print("Full Word report generated successfully!")
